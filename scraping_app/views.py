from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib import messages
from .forms import UrlForm
from scraping_tools.scraping import ejecutar_scraping  # Función que llamará al script de scraping
from django.http import HttpResponse
import pandas as pd
import csv
import io
from django.contrib.auth.decorators import login_required
import os
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait 
from selenium.webdriver.support import expected_conditions as EC
import time
from django.http import JsonResponse
import threading
from selenium.webdriver.common.by import By
import mysql.connector
from mysql.connector import Error
from datetime import datetime

def guardar_coincidencias_bd(coincidencias_df, post_url):
    try:
        conexion = mysql.connector.connect(
            host='158.69.26.160',
            user='admin',
            password='S3gur1d4d2025',
            database='fb_scrap'
        )
        cursor = conexion.cursor()

        coincidencias_df['perfil'] = coincidencias_df['perfil'].str.split('?').str[0].str.strip()

        for _, fila in coincidencias_df.iterrows():
            perfil = fila['perfil']         # URL del usuario (facebook_user)
            texto = fila['texto']           # Comentario
            nombre = fila['nombre']         # Nombre del usuario (usuario)

            # 🔁 CORREGIDO: buscar por facebook_user, no por "perfil"
            cursor.execute("SELECT id FROM empleados WHERE facebook_user = %s", (perfil,))
            resultado_empleado = cursor.fetchone()
            if not resultado_empleado:
                continue  # No está registrado, lo omitimos
            empleado_id = resultado_empleado[0]

            # Obtener ID del comentario por usuario, comentario y post_url
            cursor.execute("""
                SELECT id FROM comentarios 
                WHERE usuario = %s AND comentario = %s AND post_url = %s
                ORDER BY fecha DESC LIMIT 1
            """, (nombre, texto, post_url))
            resultado_comentario = cursor.fetchone()
            if not resultado_comentario:
                continue
            comentario_id = resultado_comentario[0]

            # Insertar coincidencia incluyendo usuario_url
            cursor.execute("""
                INSERT INTO coincidencias (post_url, usuario_url, empleado_id, comentario_id, fecha)
                VALUES (%s, %s, %s, %s, NOW())
            """, (post_url, perfil, empleado_id, comentario_id))

        conexion.commit()
        print(f"{cursor.rowcount} coincidencias insertadas.")
    except Exception as e:
        print(f"Error al insertar coincidencias: {str(e)}")
    finally:
        if conexion.is_connected():
            cursor.close()
            conexion.close()

@login_required(login_url='/login/')
def index(request):
    if request.method == 'POST':
        form = UrlForm(request.POST)
        if form.is_valid():
            url = form.cleaned_data['url']
            comentarios = ejecutar_scraping(url)

            try:
                empleados_df = pd.read_csv('scraping_tools/empleados.csv')
                comentarios_df = pd.DataFrame(comentarios)
                comentarios_df['perfil_base'] = comentarios_df['perfil'].str.extract(r'(https://www\.facebook\.com/[^/?]+)')
                empleados_df['perfil'] = empleados_df['perfil'].str.strip()
                coincidencias = comentarios_df[comentarios_df['perfil_base'].isin(empleados_df['perfil'])]
                data = coincidencias[['nombre', 'texto', 'perfil']].to_dict(orient='records')
                guardar_coincidencias_bd(coincidencias, url)
            except Exception as e:
                data = []
                print(f"Error en comparación: {e}")

            # Guardar temporalmente en la sesión (solo campos simples)
            request.session['comentarios'] = comentarios
            request.session['url_publicacion'] = url  # ← Guardar URL de publicación

            return render(request, 'scraping_app/index.html', {
                'form': form,
                'comentarios': comentarios,
                'data': data, # <- IMPORTANTE mandar también data
                'url_publicacion': url  # ← Añade esto
            })
    else:
        form = UrlForm()
    
    return render(request, 'scraping_app/index.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            return redirect('dashboard')
        else:
            messages.error(request, 'Usuario o contraseña incorrectos')
    return render(request, 'scraping_app/login.html')

def logout_view(request):
    logout(request)
    return redirect('login')

def register_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = User.objects.create_user(username=username, password=password)
        login(request, user)
        return redirect('index')
    return render(request, 'scraping_app/register.html')

def exportar_csv(request):
    comentarios = request.session.get('comentarios', [])

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="comentarios.csv"'

    writer = csv.writer(response)
    writer.writerow(['Nombre', 'Comentario', 'Perfil'])
    for c in comentarios:
        writer.writerow([c['nombre'], c['texto'], c['perfil']])
    
    return response

def exportar_excel(request):
    comentarios = request.session.get('comentarios', [])
    df = pd.DataFrame(comentarios)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Comentarios')
    
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="comentarios.xlsx"'
    return response

# Variable global para mantener la instancia del driver
driver_global = None
driver_lock = threading.Lock()

def buscar_comentario(request):
    if request.method == 'POST':
        try:
            # Validación más robusta
            texto_comentario = request.POST.get('texto', '').strip()
            url_publicacion = request.POST.get('url_publicacion', '').strip()  # Cambiado de 'url' a 'url_publicacion'
            
            if not texto_comentario or not url_publicacion:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Texto y URL son requeridos'
                }, status=400)

            global driver_global
            
            with driver_lock:
                if driver_global is None:
                    options = Options()
                    options.add_argument("--headless")  # Recomendado para producción
                    options.add_argument("--disable-notifications")
                    options.add_argument("--disable-infobars")
                    driver_global = webdriver.Chrome(options=options)
                    
                    try:
                        driver_global.get("https://www.facebook.com/")
                        # [Mantener tu código de login...]
                    except Exception as e:
                        driver_global.quit()
                        driver_global = None
                        raise

                try:
                    # Validar URL antes de usarla
                    if not url_publicacion.startswith('http'):
                        url_publicacion = f'https://{url_publicacion}'
                    
                    driver_global.get(url_publicacion)
                    WebDriverWait(driver_global, 15).until(
                        EC.presence_of_element_located((By.XPATH, '//div[contains(@role, "article")]'))
                    )
                    
                    # Buscar con esperas explícitas
                    comentarios = WebDriverWait(driver_global, 15).until(
                        EC.presence_of_all_elements_located((By.XPATH, '//div[contains(@role, "article")]'))
                    )
                    
                    for comentario in comentarios:
                        try:
                            texto = WebDriverWait(comentario, 5).until(
                                EC.presence_of_element_located((By.XPATH, './/div[@dir="auto"]'))
                            ).text
                            
                            if texto_comentario.lower() in texto.lower():
                                # Screenshot con nombre único
                                screenshot_path = f"comentario_{int(time.time())}.png"
                                comentario.screenshot(screenshot_path)
                                
                                return JsonResponse({
                                    'status': 'success',
                                    'screenshot': f'/media/{screenshot_path}'  # Asegúrate de configurar MEDIA_URL
                                })
                                
                        except Exception as e:
                            continue
                    
                    return JsonResponse({
                        'status': 'error', 
                        'message': 'Comentario no encontrado'
                    }, status=404)
                    
                except Exception as e:
                    print(f"Error durante la búsqueda: {str(e)}")
                    return JsonResponse({
                        'status': 'error',
                        'message': f'Error técnico: {str(e)}'
                    }, status=500)
                    
        except Exception as e:
            print(f"Error general: {str(e)}")
            return JsonResponse({
                'status': 'error',
                'message': 'Error interno del servidor'
            }, status=500)
    
    return JsonResponse({
        'status': 'error',
        'message': 'Método no permitido'
    }, status=405)
    
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from django.http import HttpResponse
from datetime import datetime

def exportar_pdf(request):
    url_publicacion = request.GET.get('post_url')
    
    coincidencias_qs = Coincidencias.objects.using('fb_scrap').select_related('comentario', 'empleado')
    if url_publicacion:
        coincidencias_qs = coincidencias_qs.filter(post_url=url_publicacion)
    
    data = [{
        'nombre': c.empleado.nombre if c.empleado else 'Desconocido',
        'texto': c.comentario.comentario if c.comentario else '',
        'perfil': c.comentario.usuario_url if c.comentario else '',
    } for c in coincidencias_qs]

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="reporte_coincidencias.pdf"'
    
    doc = SimpleDocTemplate(response, pagesize=letter)
    styles = getSampleStyleSheet()

    if 'TitleCustom' not in styles:
        styles.add(ParagraphStyle(name='TitleCustom', parent=styles['Heading1'], fontSize=16, alignment=TA_CENTER, spaceAfter=20, textColor=colors.darkblue))

    if 'Subtitle' not in styles:
        styles.add(ParagraphStyle(name='Subtitle', fontSize=12, alignment=TA_CENTER, spaceAfter=20, textColor=colors.grey))

    if 'ItemTitle' not in styles:
        styles.add(ParagraphStyle(name='ItemTitle', fontSize=12, textColor=colors.darkblue, spaceAfter=5))

    if 'ItemContent' not in styles:
        styles.add(ParagraphStyle(name='ItemContent', fontSize=10, spaceAfter=10, leading=14))
    
    elements = []

    # Encabezado
    elements.append(Paragraph("Reporte de Coincidencias de Perfiles", styles['TitleCustom']))
    elements.append(Paragraph(f"Publicación origen: {url_publicacion or 'No especificada'}", styles['Subtitle']))
    elements.append(Paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Subtitle']))
    elements.append(Spacer(1, 20))

    # Contenido
    if not data:
        elements.append(Paragraph("No se encontraron coincidencias para esta publicación.", styles['ItemContent']))
    else:
        for i, item in enumerate(data, start=1):
            elements.append(Paragraph(f"{i}. {item.get('nombre', 'N/A')}", styles['ItemTitle']))
            comment = item.get('texto', 'N/A').replace('\n', '<br/>')
            elements.append(Paragraph(f"<b>Comentario:</b> {comment}", styles['ItemContent']))
            elements.append(Paragraph(f"<b>Enlace al perfil:</b> {item.get('perfil', 'N/A')}", styles['ItemContent']))
            elements.append(Spacer(1, 15))

    doc.build(elements)
    return response


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def exportar_word(request):
    url_publicacion = request.GET.get('post_url')

    coincidencias_qs = Coincidencias.objects.using('fb_scrap').select_related('comentario', 'empleado')
    if url_publicacion:
        coincidencias_qs = coincidencias_qs.filter(post_url=url_publicacion)

    data = [{
        'nombre': c.empleado.nombre if c.empleado else 'Desconocido',
        'texto': c.comentario.comentario if c.comentario else '',
        'perfil': c.comentario.usuario_url if c.comentario else '',
    } for c in coincidencias_qs]

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="reporte_coincidencias.docx"'
    
    document = Document()
    
    # Estilo para el título
    title = document.add_heading('Reporte de Coincidencias de Perfiles', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Agregar URL de publicación
    document.add_paragraph(f"Publicación origen: {url_publicacion}", style='Intense Quote')
    
    # Estilo para el contenido
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Agregar cada coincidencia con formato legible
    for i, item in enumerate(data, start=1):
        # Nombre
        p = document.add_paragraph()
        p.add_run(f"{i}. Nombre: ").bold = True
        p.add_run(f"{item.get('nombre', 'N/A')}")
        
        # Comentario
        p = document.add_paragraph()
        p.add_run("Comentario: ").bold = True
        p.add_run(f"{item.get('texto', 'N/A')}")
        
        # Perfil
        p = document.add_paragraph()
        p.add_run("Enlace al perfil: ").bold = True
        p.add_run(f"{item.get('perfil', 'N/A')}")
        
        # Espaciado entre registros
        document.add_paragraph()
    
    document.save(response)
    return response

#--------------------------------------------------------------------------------#
# Implementación del Dashboard y nueva vista principal
from django.db.models import Q
from django.shortcuts import render
from datetime import datetime
from .models import Comentarios, Coincidencias, Empleados, Clientes

def dashboard(request):
    desde_str = request.GET.get('desde')
    hasta_str = request.GET.get('hasta')
    cliente_id = request.GET.get('cliente')
    palabra_clave = request.GET.get('palabra_clave')
    empleado_id = request.GET.get('empleado')

    filtros = {}
    comentarios_qs = Comentarios.objects.using('fb_scrap').all()
    coincidencias_qs = Coincidencias.objects.using('fb_scrap').all()

    # Filtro fechas
    if desde_str:
        comentarios_qs = comentarios_qs.filter(fecha__gte=desde_str)
        coincidencias_qs = coincidencias_qs.filter(fecha__gte=desde_str)
        filtros['desde'] = desde_str
    if hasta_str:
        comentarios_qs = comentarios_qs.filter(fecha__lte=hasta_str)
        coincidencias_qs = coincidencias_qs.filter(fecha__lte=hasta_str)
        filtros['hasta'] = hasta_str

    # Filtro cliente (solo si usuario staff y cliente seleccionado)
    if request.user.is_staff and cliente_id:
        nombre_cliente = Clientes.objects.filter(id=cliente_id).values_list('nombre_pagina', flat=True).first()
        if nombre_cliente:
            comentarios_qs = comentarios_qs.filter(nombre_pagina=nombre_cliente)
            coincidencias_qs = coincidencias_qs.filter(comentario__nombre_pagina=nombre_cliente)
            filtros['cliente_id'] = int(cliente_id)

    # Filtro palabra clave (busca texto en comentarios)
    if palabra_clave:
        comentarios_qs = comentarios_qs.filter(comentario__icontains=palabra_clave)
        coincidencias_qs = coincidencias_qs.filter(comentario__comentario__icontains=palabra_clave)
        filtros['palabra_clave'] = palabra_clave

    # Filtro empleado (por coincidencias)
    if empleado_id:
        coincidencias_qs = coincidencias_qs.filter(empleado__id=empleado_id)
        filtros['empleado_id'] = int(empleado_id)

    # Métricas
    total_comentarios = comentarios_qs.count()
    total_coincidencias = coincidencias_qs.count()
    total_posts = comentarios_qs.values('post_url').distinct().count()
    total_no_empleados = total_comentarios - total_coincidencias
    porcentaje_coincidencias = round((total_coincidencias / total_comentarios) * 100, 1) if total_comentarios else 0

    # Comentarios y coincidencias recientes
    comentarios_data = comentarios_qs.order_by('-fecha')
    coincidencias_data = coincidencias_qs.order_by('-fecha')

    from django.utils.timezone import now, timedelta

    # Rango de fechas
    hoy = now().date()
    ultimos_7_dias = [hoy - timedelta(days=i) for i in range(6, -1, -1)]

    # Crear listas de fechas y cantidades
    coincidencias_labels = []
    coincidencias_por_dia = []

    for dia in ultimos_7_dias:
        coincidencias_labels.append(dia.strftime('%Y-%m-%d'))
        coincidencias_por_dia.append(
            coincidencias_qs.filter(fecha__date=dia).count()
        )

    comentarios = [{
        'usuario_nombre': c.usuario,
        'usuario_url': c.usuario_url,
        'texto': c.comentario,
        'fecha': c.fecha,
        'post_url': c.post_url,
    } for c in comentarios_data]

    coincidencias = [{
        'empleado_nombre': c.empleado.nombre if c.empleado else 'Desconocido',
        'usuario_url': c.comentario.usuario_url if c.comentario else 'Desconocido',
        'texto': c.comentario.comentario if c.comentario else 'Comentario no disponible',
        'fecha': c.fecha,
        'post_url': c.post_url,
    } for c in coincidencias_data]

    context = {
        'filtros': filtros,
        'is_staff': request.user.is_staff,
        'clientes': Clientes.objects.using('fb_scrap').all() if request.user.is_staff else [],
        'empleados': Empleados.objects.using('fb_scrap').all(),
        'no_empleados': total_no_empleados,
        'total_comentarios': total_comentarios,
        'total_coincidencias': total_coincidencias,
        'total_posts': total_posts,
        'porcentaje_coincidencias': porcentaje_coincidencias,
        'comentarios': comentarios,
        'coincidencias': coincidencias,
        'coincidencias_labels': coincidencias_labels,
        'coincidencias_data': coincidencias_por_dia,
    }

    return render(request, 'dashboard.html', context)
